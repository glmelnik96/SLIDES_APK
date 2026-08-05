"""DOCX-парсер: структура из стилей и — если стилей нет — из оформления.

Реальные клиентские .docx сплошь и рядом свёрстаны руками: ни одного стиля
Heading, заголовки набраны жирным/капсом, списки — символом «·» с клавиатуры.
Парсер отдавал такой документ одной секцией (см.
docs/superpowers/specs/2026-08-05-docx-structure-heuristics.md), поэтому здесь
проверяется и стилевой путь (он не должен измениться), и эвристический.
"""
from __future__ import annotations

import docx as docx_lib
import pytest
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from htmlslides.parsers.base import ListBlock, TextBlock
from htmlslides.parsers.docx import parse_docx


def _parse(tmp_path, build):
    document = docx_lib.Document()
    build(document)
    path = tmp_path / "case.docx"
    document.save(str(path))
    return parse_docx(path)


def _bold(document, text):
    paragraph = document.add_paragraph()
    paragraph.add_run(text).bold = True
    return paragraph


def _texts(section):
    return [b.text for b in section.blocks if isinstance(b, TextBlock)]


# --- стилевой путь: поведение не меняется -----------------------------------

def test_styled_headings_still_build_sections(tmp_path):
    def build(d):
        d.add_paragraph("Заголовок деки", style="Title")
        d.add_paragraph("Первый", style="Heading 1")
        d.add_paragraph("тело первого")
        d.add_paragraph("Второй", style="Heading 2")
        d.add_paragraph("тело второго")

    doc = _parse(tmp_path, build)
    assert doc.title == "Заголовок деки"
    assert [(s.heading, s.level) for s in doc.sections] == [("Первый", 1), ("Второй", 2)]


def test_styles_present_disable_heuristics(tmp_path):
    """Документ размечен стилями → жирный абзац остаётся абзацем.

    Иначе на нормальных документах жирный акцент внутри текста разорвал бы
    раздел пополам — регрессия там, где сегодня всё хорошо.
    """
    def build(d):
        d.add_paragraph("Первый", style="Heading 1")
        _bold(d, "ВАЖНАЯ МЫСЛЬ")
        d.add_paragraph("тело")

    doc = _parse(tmp_path, build)
    assert len(doc.sections) == 1
    assert _texts(doc.sections[0]) == ["ВАЖНАЯ МЫСЛЬ", "тело"]


def test_outline_level_disables_heuristics(tmp_path):
    """Структура может быть задана не стилем, а w:outlineLvl — это тоже разметка."""
    def build(d):
        p = d.add_paragraph("Раздел через outlineLvl")
        lvl = OxmlElement("w:outlineLvl")
        lvl.set(qn("w:val"), "0")
        p._p.get_or_add_pPr().append(lvl)
        _bold(d, "ЖИРНЫЙ АБЗАЦ")

    doc = _parse(tmp_path, build)
    assert len(doc.sections) == 1
    assert _texts(doc.sections[0]) == ["Раздел через outlineLvl", "ЖИРНЫЙ АБЗАЦ"]


# --- эвристический путь: заголовки ------------------------------------------

def test_bold_line_becomes_heading(tmp_path):
    def build(d):
        _bold(d, "Общая концепция")
        d.add_paragraph("тело первого")
        _bold(d, "Хронометраж")
        d.add_paragraph("тело второго")

    doc = _parse(tmp_path, build)
    assert [(s.heading, s.level) for s in doc.sections] == [
        ("Общая концепция", 1), ("Хронометраж", 1)]
    assert _texts(doc.sections[0]) == ["тело первого"]


def test_caps_line_becomes_heading(tmp_path):
    def build(d):
        d.add_paragraph("2. ХРОНОМЕТРАЖ И ФОРМАТ")
        d.add_paragraph("тело")

    doc = _parse(tmp_path, build)
    assert [s.heading for s in doc.sections] == ["2. ХРОНОМЕТРАЖ И ФОРМАТ"]


def test_text_before_first_heading_kept_as_preamble(tmp_path):
    """Вступление до первого заголовка терять нельзя — там часто вся задача."""
    def build(d):
        d.add_paragraph("Вступительный абзац без заголовка.")
        _bold(d, "Раздел")
        d.add_paragraph("тело")

    doc = _parse(tmp_path, build)
    assert doc.sections[0].heading == ""
    assert _texts(doc.sections[0]) == ["Вступительный абзац без заголовка."]
    assert doc.sections[1].heading == "Раздел"


@pytest.mark.parametrize("text", [
    "Это обычное предложение, набранное жирным.",     # точка в конце
    "Приоритеты при любом конфликте:",                # двоеточие = дальше текст
    "Ж" * 91,                                          # длинная строка
])
def test_not_a_heading(tmp_path, text):
    def build(d):
        _bold(d, text)
        d.add_paragraph("тело")

    doc = _parse(tmp_path, build)
    assert doc.sections[0].heading == ""
    assert text in _texts(doc.sections[0])


def test_partially_bold_line_is_not_heading(tmp_path):
    """Жирное только первое слово — это акцент внутри абзаца, а не заголовок."""
    def build(d):
        p = d.add_paragraph()
        p.add_run("ВАЖНО").bold = True
        p.add_run(" остальное обычным")
        d.add_paragraph("тело")

    doc = _parse(tmp_path, build)
    assert doc.sections[0].heading == ""


def test_digits_alone_are_not_caps_heading(tmp_path):
    """«2024» равно себе в верхнем регистре — но это не заголовок."""
    def build(d):
        d.add_paragraph("2024")
        d.add_paragraph("тело")

    doc = _parse(tmp_path, build)
    assert doc.sections[0].heading == ""
    assert _texts(doc.sections[0]) == ["2024", "тело"]


# --- эвристический путь: списки ---------------------------------------------

def test_literal_bullets_become_list(tmp_path):
    def build(d):
        _bold(d, "Раздел")
        d.add_paragraph("· первый пункт")
        d.add_paragraph("• второй пункт")
        d.add_paragraph("- третий пункт")
        d.add_paragraph("обычный абзац")

    doc = _parse(tmp_path, build)
    blocks = doc.sections[0].blocks
    assert isinstance(blocks[0], ListBlock)
    assert blocks[0].items == ["первый пункт", "второй пункт", "третий пункт"]
    assert blocks[0].ordered is False
    assert _texts(doc.sections[0]) == ["обычный абзац"]


def test_bullet_wins_over_heading(tmp_path):
    """Пункт списка капсом остаётся пунктом: список важнее оформления."""
    def build(d):
        _bold(d, "Раздел")
        d.add_paragraph("· ПЕРВЫЙ ПУНКТ")
        d.add_paragraph("· ВТОРОЙ ПУНКТ")

    doc = _parse(tmp_path, build)
    assert len(doc.sections) == 1
    assert doc.sections[0].blocks[0].items == ["ПЕРВЫЙ ПУНКТ", "ВТОРОЙ ПУНКТ"]


def test_caps_item_of_word_list_is_heading(tmp_path):
    """Автонумерация Word на первом заголовке — реальный кейс: дальше номера руками."""
    def build(d):
        d.add_paragraph("ОБЩАЯ КОНЦЕПЦИЯ", style="List Number")
        d.add_paragraph("тело")

    doc = _parse(tmp_path, build)
    assert [s.heading for s in doc.sections] == ["ОБЩАЯ КОНЦЕПЦИЯ"]


def test_bold_item_of_word_list_stays_item(tmp_path):
    """Внутри списка жирности мало: жирный пункт — обычное оформление."""
    def build(d):
        _bold(d, "Раздел")
        p = d.add_paragraph(style="List Bullet")
        p.add_run("Первый пункт").bold = True

    doc = _parse(tmp_path, build)
    assert len(doc.sections) == 1
    assert doc.sections[0].blocks[0].items == ["Первый пункт"]


def test_dash_without_space_is_not_a_bullet(tmp_path):
    """«—» внутри фразы не должен превращать абзац в список."""
    def build(d):
        _bold(d, "Раздел")
        d.add_paragraph("-30% к бюджету")

    doc = _parse(tmp_path, build)
    assert _texts(doc.sections[0]) == ["-30% к бюджету"]
